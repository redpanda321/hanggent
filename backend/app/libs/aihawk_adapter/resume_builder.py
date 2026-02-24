"""
Resume Builder using AIHawk's resume generation capabilities.

This module wraps AIHawk's ResumeFacade and ResumeGenerator to work with
Hanggent's model configuration and file storage patterns.

Supports two input modes:
1. Dict mode (legacy): Pass resume_data dict for backwards compatibility
2. YAML mode (preferred): Pass resume_yaml string from JobAnalyzerToolkit.extract_resume_to_yaml()
   for full structured data support
"""
import os
import sys
import hashlib
from pathlib import Path
from typing import Optional, Dict, List, Tuple, Any

from app.model.chat import Chat
from app.libs.aihawk_adapter.llm_adapter import inject_llm_into_aihawk
from utils import traceroot_wrapper as traceroot

logger = traceroot.get_logger("aihawk_resume_builder")

# Add AIHawk to Python path
AIHAWK_PATH = Path(__file__).resolve().parents[4] / "external" / "aihawk"
AIHAWK_SRC_PATH = AIHAWK_PATH / "src"

if str(AIHAWK_SRC_PATH) not in sys.path:
    sys.path.insert(0, str(AIHAWK_SRC_PATH))
if str(AIHAWK_PATH) not in sys.path:
    sys.path.insert(0, str(AIHAWK_PATH))


class AIHawkResumeBuilder:
    """
    High-level interface for AIHawk resume generation.
    
    Wraps AIHawk's ResumeFacade with Hanggent's configuration and
    provides simplified methods for resume tailoring.
    
    Supports two input modes:
    - resume_data (dict): Legacy mode with parsed resume fields
    - resume_yaml (str): Preferred mode with AIHawk-compatible YAML
    """
    
    def __init__(
        self,
        chat_options: Chat,
        output_directory: str,
        resume_data: Optional[Dict] = None,
        resume_yaml: Optional[str] = None,
    ):
        """
        Initialize the resume builder.
        
        Args:
            chat_options: Hanggent Chat options with LLM configuration
            output_directory: Directory for generated resume files
            resume_data: Parsed resume data dict (legacy mode)
            resume_yaml: YAML string from extract_resume_to_yaml (preferred mode)
        """
        self.chat_options = chat_options
        self.output_directory = Path(output_directory)
        self.output_directory.mkdir(parents=True, exist_ok=True)
        self.resume_data = resume_data or {}
        self.resume_yaml = resume_yaml
        
        # Inject API key into environment for AIHawk
        self.api_key = inject_llm_into_aihawk(chat_options)
        
        # Lazy-loaded AIHawk components
        self._style_manager = None
        self._resume_generator = None
        self._facade = None
        self._resume_object = None
        
    def _get_aihawk_components(self):
        """Lazy-load AIHawk components"""
        if self._facade is not None:
            return
            
        try:
            from src.libs.resume_and_cover_builder.resume_facade import ResumeFacade
            from src.libs.resume_and_cover_builder.resume_generator import ResumeGenerator
            from src.libs.resume_and_cover_builder.style_manager import StyleManager
            from src.libs.resume_and_cover_builder.config import global_config
            
            # Set up global config
            lib_dir = AIHAWK_SRC_PATH / "libs" / "resume_and_cover_builder"
            global_config.STYLES_DIRECTORY = lib_dir / "resume_style"
            global_config.API_KEY = self.api_key
            
            # Initialize components
            self._style_manager = StyleManager()
            self._resume_generator = ResumeGenerator()
            
            # Create resume object (prefer YAML mode)
            self._resume_object = self._create_resume_object()
            
            # Create facade
            self._facade = ResumeFacade(
                api_key=self.api_key,
                style_manager=self._style_manager,
                resume_generator=self._resume_generator,
                resume_object=self._resume_object,
                output_path=str(self.output_directory),
            )
            
            logger.info("AIHawk components initialized successfully")
            
        except ImportError as e:
            logger.error(f"Failed to import AIHawk components: {e}")
            raise ImportError(
                "AIHawk components not found. Make sure the submodule is initialized: "
                "git submodule update --init --recursive"
            ) from e
    
    def _create_resume_object(self) -> Any:
        """
        Convert resume input to AIHawk Resume schema.
        
        Supports two input modes:
        1. YAML mode (preferred): Uses resume_yaml string directly
        2. Dict mode (legacy): Maps resume_data dict to AIHawk schema
        
        Returns:
            Resume object compatible with AIHawk
        """
        try:
            from src.resume_schemas.resume import Resume, PersonalInformation, EducationDetails, ExperienceDetails
            
            # YAML mode (preferred) - construct Resume directly from YAML
            if self.resume_yaml:
                logger.info("Creating Resume object from YAML input")
                try:
                    # AIHawk's Resume class accepts yaml_str in constructor
                    return Resume(yaml_str=self.resume_yaml)
                except TypeError:
                    # Fallback: parse YAML and pass as dict
                    import yaml
                    yaml_data = yaml.safe_load(self.resume_yaml)
                    return self._create_resume_from_dict(yaml_data, Resume, PersonalInformation, EducationDetails, ExperienceDetails)
            
            # Legacy dict mode - map Hanggent resume_data to AIHawk schema
            return self._create_resume_from_legacy_dict(Resume, PersonalInformation, EducationDetails, ExperienceDetails)
            
        except ImportError:
            logger.warning("AIHawk Resume schema not available, using dict")
            return self.resume_data if not self.resume_yaml else {}
    
    def _create_resume_from_dict(
        self,
        data: Dict,
        Resume,
        PersonalInformation,
        EducationDetails,
        ExperienceDetails,
    ) -> Any:
        """Create Resume object from plain_text_resume YAML structure."""
        pi_data = data.get("personal_information", {})
        personal_info = PersonalInformation(
            name=pi_data.get("name", ""),
            surname=pi_data.get("surname", ""),
            email=pi_data.get("email", ""),
            phone=pi_data.get("phone", ""),
            github=pi_data.get("github", ""),
            linkedin=pi_data.get("linkedin", ""),
        ) if pi_data else None
        
        education_list = []
        for edu in data.get("education_details", []):
            education_list.append(EducationDetails(
                education_level=edu.get("education_level", ""),
                institution=edu.get("institution", ""),
                field_of_study=edu.get("field_of_study", ""),
                final_evaluation_grade=edu.get("final_evaluation_grade", ""),
                start_date=edu.get("start_date", ""),
                year_of_completion=edu.get("year_of_completion", ""),
            ))
        
        experience_list = []
        for exp in data.get("experience_details", []):
            experience_list.append(ExperienceDetails(
                position=exp.get("position", ""),
                company=exp.get("company", ""),
                employment_period=exp.get("employment_period", ""),
                location=exp.get("location", ""),
                industry=exp.get("industry", ""),
                key_responsibilities=exp.get("key_responsibilities", []),
                skills_acquired=exp.get("skills_acquired", []),
            ))
        
        return Resume(
            personal_information=personal_info,
            education_details=education_list if education_list else None,
            experience_details=experience_list if experience_list else None,
            projects=data.get("projects"),
            achievements=data.get("achievements"),
            certifications=data.get("certifications"),
            languages=data.get("languages"),
            interests=data.get("interests", []),
        )
    
    def _create_resume_from_legacy_dict(
        self,
        Resume,
        PersonalInformation,
        EducationDetails,
        ExperienceDetails,
    ) -> Any:
        """Create Resume from legacy Hanggent resume_data dict format."""
        # Map Hanggent resume data to AIHawk schema
        personal_info = PersonalInformation(
            name=self.resume_data.get("name", ""),
            surname="",  # AIHawk separates first/last name
            email=self.resume_data.get("email", ""),
            phone=self.resume_data.get("phone", ""),
            github="",
            linkedin="",
        ) if self.resume_data.get("name") else None
        
        education_list = []
        for edu in self.resume_data.get("education", []):
            education_list.append(EducationDetails(
                education_level=edu.get("degree", ""),
                institution=edu.get("school", ""),
                field_of_study=edu.get("field", ""),
                final_evaluation_grade=edu.get("gpa", ""),
                start_date="",
                year_of_completion=edu.get("year", ""),
            ))
        
        experience_list = []
        for exp in self.resume_data.get("experience", []):
            experience_list.append(ExperienceDetails(
                position=exp.get("title", ""),
                company=exp.get("company", ""),
                employment_period=exp.get("duration", ""),
                location="",
                industry="",
                key_responsibilities=exp.get("bullets", []),
                skills_acquired=exp.get("skills", []),
            ))
        
        return Resume(
            personal_information=personal_info,
            education_details=education_list if education_list else None,
            experience_details=experience_list if experience_list else None,
            projects=None,
            achievements=None,
            certifications=None,
            languages=None,
            interests=self.resume_data.get("skills", []),
        )
    
    def tailor_resume_for_job(
        self,
        job_description: str,
        style: str = "classic",
    ) -> Tuple[str, str]:
        """
        Generate a tailored resume for a specific job description.
        
        Args:
            job_description: Full job description text
            style: Resume style (classic, modern, etc.)
            
        Returns:
            Tuple of (html_content, markdown_content)
        """
        self._get_aihawk_components()
        
        try:
            # Set style
            style_path = self._get_style_path(style)
            
            # Generate tailored resume HTML
            html_content = self._resume_generator.create_resume_job_description_text(
                style_path=str(style_path),
                job_description_text=job_description,
            )
            
            # Convert HTML to markdown (simplified)
            markdown_content = self._html_to_markdown(html_content)
            
            logger.info(f"Generated tailored resume with style: {style}")
            return html_content, markdown_content
            
        except Exception as e:
            logger.error(f"Error generating tailored resume: {e}")
            raise
    
    def generate_cover_letter(
        self,
        job_description: str,
        style: str = "classic",
    ) -> Tuple[str, str]:
        """
        Generate a tailored cover letter for a specific job.
        
        Args:
            job_description: Full job description text
            style: Cover letter style
            
        Returns:
            Tuple of (html_content, markdown_content)
        """
        self._get_aihawk_components()
        
        try:
            style_path = self._get_style_path(style)
            
            html_content = self._resume_generator.create_cover_letter_job_description(
                style_path=str(style_path),
                job_description_text=job_description,
            )
            
            markdown_content = self._html_to_markdown(html_content)
            
            logger.info("Generated cover letter")
            return html_content, markdown_content
            
        except Exception as e:
            logger.error(f"Error generating cover letter: {e}")
            raise
    
    def export_to_pdf(
        self,
        html_content: str,
        filename: str,
    ) -> str:
        """
        Export HTML content to PDF.
        
        Args:
            html_content: HTML resume/cover letter content
            filename: Output filename (without extension)
            
        Returns:
            Path to the generated PDF file
        """
        output_path = self.output_directory / f"{filename}.pdf"
        
        try:
            import pdfkit
            pdfkit.from_string(
                html_content,
                str(output_path),
                options={
                    'page-size': 'Letter',
                    'margin-top': '0.5in',
                    'margin-right': '0.5in',
                    'margin-bottom': '0.5in',
                    'margin-left': '0.5in',
                    'encoding': 'UTF-8',
                }
            )
            logger.info(f"Exported PDF to: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            # Fallback: save HTML
            html_path = self.output_directory / f"{filename}.html"
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            logger.info(f"Saved HTML fallback to: {html_path}")
            return str(html_path)
    
    def export_to_docx(
        self,
        markdown_content: str,
        filename: str,
    ) -> str:
        """
        Export markdown content to DOCX.
        
        Args:
            markdown_content: Markdown resume content
            filename: Output filename (without extension)
            
        Returns:
            Path to the generated DOCX file
        """
        output_path = self.output_directory / f"{filename}.docx"
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            
            doc = Document()
            
            # Parse markdown and add to document
            lines = markdown_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)
            
            doc.save(str(output_path))
            logger.info(f"Exported DOCX to: {output_path}")
            return str(output_path)
            
        except ImportError:
            logger.error("python-docx not installed")
            # Fallback: save as markdown
            md_path = self.output_directory / f"{filename}.md"
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            return str(md_path)
    
    def _get_style_path(self, style: str) -> Path:
        """Get the path to a resume style CSS file"""
        lib_dir = AIHAWK_SRC_PATH / "libs" / "resume_and_cover_builder"
        styles_dir = lib_dir / "resume_style"
        
        # Map style names to files
        style_files = {
            "classic": "classic.css",
            "modern": "modern.css",
            "minimal": "minimal.css",
            "professional": "professional.css",
        }
        
        style_file = style_files.get(style.lower(), "classic.css")
        style_path = styles_dir / style_file
        
        if not style_path.exists():
            # Fallback to first available style
            for css_file in styles_dir.glob("*.css"):
                return css_file
            raise FileNotFoundError(f"No style files found in {styles_dir}")
        
        return style_path
    
    def _html_to_markdown(self, html_content: str) -> str:
        """Convert HTML to markdown (simplified conversion)"""
        import re
        
        # Basic HTML to markdown conversion
        text = html_content
        
        # Remove HTML tags but preserve content structure
        text = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1\n', text, flags=re.DOTALL)
        text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1\n', text, flags=re.DOTALL)
        text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1\n', text, flags=re.DOTALL)
        text = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', text, flags=re.DOTALL)
        text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text, flags=re.DOTALL)
        text = re.sub(r'<br\s*/?>', '\n', text)
        text = re.sub(r'<[^>]+>', '', text)  # Remove remaining tags
        text = re.sub(r'\n{3,}', '\n\n', text)  # Clean up extra newlines
        
        return text.strip()
    
    def get_available_styles(self) -> List[str]:
        """Get list of available resume styles"""
        lib_dir = AIHAWK_SRC_PATH / "libs" / "resume_and_cover_builder"
        styles_dir = lib_dir / "resume_style"
        
        styles = []
        if styles_dir.exists():
            for css_file in styles_dir.glob("*.css"):
                styles.append(css_file.stem)
        
        return styles if styles else ["classic"]
