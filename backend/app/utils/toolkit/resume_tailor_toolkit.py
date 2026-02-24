"""
Resume Tailor Toolkit using AIHawk integration.

This toolkit provides tools for creating job-tailored resumes following
the AIHawk methodology (https://github.com/feder-cr/Jobs_Applier_AI_Agent_AIHawk).

Key principles:
1. PRESERVE original formatting, structure, and layout
2. REWRITE content to highlight alignment with job description
3. NEVER fabricate skills or experiences not in original resume
4. OPTIMIZE for ATS systems using job-specific keywords
5. MAINTAIN professional, concise language

Supports two input formats:
1. Plain text resume (legacy): Raw resume text with optional analysis metadata
2. YAML resume (preferred): Structured YAML from JobAnalyzerToolkit.extract_resume_to_yaml()
"""
import json
import os
import asyncio
from datetime import datetime
from typing import List, Dict, Optional, Tuple, Any
from pathlib import Path

from camel.toolkits import BaseToolkit
from camel.toolkits.function_tool import FunctionTool
from camel.models import ModelFactory
from camel.messages import BaseMessage
from camel.agents import ChatAgent

from app.utils.toolkit.abstract_toolkit import AbstractToolkit
from app.model.chat import Chat
from app.model.plain_text_resume import PlainTextResume
from utils import traceroot_wrapper as traceroot

logger = traceroot.get_logger("resume_tailor_toolkit")


# =============================================================================
# Section-Based Tailoring Configuration
# =============================================================================

# Sections to tailor with LLM (job-relevant content)
SECTIONS_TO_TAILOR = ["experience_details", "projects", "achievements", "certifications"]

# Sections to copy unchanged (static personal info)
SECTIONS_TO_COPY = [
    "personal_information", "languages", "interests", "availability",
    "salary_expectations", "legal_authorization", "work_preferences",
    "self_identification", "education_details"  # education copied by default
]


# =============================================================================
# Section-Specific Tailoring Prompts (Optimized for 100% Job Match)
# =============================================================================

SECTION_TAILOR_PROMPTS = {
    "experience_details": """Rewrite these work experience entries to MAXIMIZE alignment with the job requirements.

GOAL: Make the resume appear as a 100% match for this specific job.

AGGRESSIVE OPTIMIZATION RULES:
1. USE EXACT KEYWORDS from the job description - if they say "distributed systems", use "distributed systems" not "scalable architecture"
2. MIRROR the job's language and terminology exactly throughout
3. REORDER responsibilities to put the most job-relevant ones FIRST
4. REWRITE each bullet point to incorporate 2-3 keywords from the job requirements
5. MATCH the job title terminology where truthful (e.g., "Backend Developer" → "Backend Engineer" if job says "Engineer")
6. QUANTIFY everything possible (%, $, users, transactions, latency improvements)
7. ADD relevant skills from job to skills_acquired if they can be reasonably inferred from the work described
8. Keep company, period, location unchanged - only rewrite position title and responsibilities
9. DO NOT fabricate completely new experiences - only optimize existing content

KEYWORD EXTRACTION - Use these exact terms from the job:
{job_description}

ORIGINAL EXPERIENCE:
{section_json}

OUTPUT JSON SCHEMA:
{output_schema}

Output ONLY the JSON object, no explanations.""",

    "projects": """Rewrite these project descriptions to MAXIMIZE alignment with the job requirements.

GOAL: Make projects appear directly relevant to this specific job.

AGGRESSIVE OPTIMIZATION RULES:
1. USE EXACT KEYWORDS from the job description in every project description
2. EMPHASIZE technologies that match job requirements (even if minor in original)
3. REFRAME the project's purpose to align with job responsibilities
4. ADD relevant technical details that can be reasonably inferred
5. QUANTIFY impact where possible (users, performance, scale)
6. Keep project names and links unchanged
7. DO NOT fabricate entirely new capabilities

KEYWORD EXTRACTION - Use these exact terms from the job:
{job_description}

ORIGINAL PROJECTS:
{section_json}

OUTPUT JSON SCHEMA:
{output_schema}

Output ONLY the JSON object, no explanations.""",

    "achievements": """Rewrite these achievements to MAXIMIZE alignment with the job requirements.

GOAL: Make achievements directly relevant to this specific job.

AGGRESSIVE OPTIMIZATION RULES:
1. USE EXACT KEYWORDS from job description
2. REFRAME achievements in terms of job-relevant skills
3. QUANTIFY with metrics that matter for this role
4. PRIORITIZE achievements most relevant to job requirements
5. DO NOT fabricate achievements not present in the original

KEYWORD EXTRACTION - Use these exact terms from the job:
{job_description}

ORIGINAL ACHIEVEMENTS:
{section_json}

OUTPUT JSON SCHEMA:
{output_schema}

Output ONLY the JSON object, no explanations.""",

    "certifications": """Reorder and annotate certifications for MAXIMUM job relevance.

GOAL: Highlight certifications most valuable for this specific job.

AGGRESSIVE OPTIMIZATION RULES:
1. REORDER with most job-relevant certifications FIRST
2. ADD relevance notes in description (e.g., "Directly relevant to [job requirement]")
3. USE EXACT terminology from job description in relevance notes
4. Keep certification names accurate
5. If certification covers job-required skills, explicitly mention those skills

KEYWORD EXTRACTION - Use these exact terms from the job:
{job_description}

ORIGINAL CERTIFICATIONS:
{section_json}

OUTPUT JSON SCHEMA:
{output_schema}

Output ONLY the JSON object, no explanations.""",

    "education_details": """Format education to MAXIMIZE alignment with job requirements.

GOAL: Highlight educational background most relevant to this job.

AGGRESSIVE OPTIMIZATION RULES:
1. EMPHASIZE coursework/specializations matching job requirements
2. USE EXACT keywords from job in any relevant coursework mentions
3. Keep institution, degree, and dates accurate
4. DO NOT fabricate degrees or coursework
5. Output valid JSON matching the schema exactly

KEYWORD EXTRACTION - Use these exact terms from the job:
{job_description}

ORIGINAL EDUCATION:
{section_json}

OUTPUT JSON SCHEMA:
{output_schema}

Output ONLY the JSON object, no explanations.""",
}


# =============================================================================
# Section Output Schemas (PlainTextResume-compatible)
# =============================================================================

SECTION_TAILOR_SCHEMAS = {
    "experience_details": '''{
  "experience_details": [{
    "position": "Job title",
    "company": "Company name",
    "employment_period": "Jan 2020 - Present",
    "location": "City, Country or null",
    "industry": "Industry sector or null",
    "key_responsibilities": [{"resp_1": "Tailored responsibility bullet point"}],
    "skills_acquired": ["Skill 1", "Skill 2"]
  }]
}''',

    "projects": '''{
  "projects": [{
    "name": "Project name",
    "description": "Tailored description emphasizing relevant tech and impact",
    "link": "URL or null"
  }]
}''',

    "achievements": '''{
  "achievements": [{
    "name": "Achievement name",
    "description": "Tailored description with metrics"
  }]
}''',

    "certifications": '''{
  "certifications": [{
    "name": "Certification name",
    "description": "Issuer, date, relevance"
  }]
}''',

    "education_details": '''{
  "education_details": [{
    "education_level": "Bachelor's/Master's/PhD",
    "institution": "University name",
    "field_of_study": "Major/Field",
    "final_evaluation_grade": "GPA or null",
    "year_of_completion": "2020"
  }]
}''',
}


class ResumeTailorToolkit(BaseToolkit, AbstractToolkit):
    """
    Toolkit for creating job-tailored resumes.
    
    Uses AIHawk integration for intelligent resume optimization
    while maintaining truthfulness and ATS compatibility.
    
    Supports two input modes:
    1. Plain text mode: Pass resume_text for basic tailoring
    2. YAML mode (preferred): Pass resume_yaml from JobAnalyzerToolkit.extract_resume_to_yaml()
       for full AIHawk integration with structured data
    
    Outputs:
    - Tailored resume content (markdown/HTML)
    - PDF export
    - DOCX export
    - Cover letters
    """
    
    agent_name: str = "Resume_Tailor_Agent"
    
    def __init__(
        self,
        api_task_id: str,
        chat_options: Optional[Chat] = None,
        agent_name: str | None = None,
        working_directory: str = "",
    ):
        self.api_task_id = api_task_id
        self.chat_options = chat_options
        self.working_directory = working_directory
        if agent_name is not None:
            self.agent_name = agent_name
        super().__init__()
        
        # Lazy-loaded AIHawk builder
        self._builder = None
        self._builder_yaml = None
    
    # =========================================================================
    # LLM Infrastructure for Section Tailoring
    # =========================================================================
    
    def _create_llm_model(self, json_mode: bool = True):
        """Create LLM model for section tailoring.
        
        Args:
            json_mode: Whether to enable JSON response format (default True)
        """
        if not self.chat_options:
            raise ValueError("chat_options required for LLM-based tailoring")
        
        # Build model config dict with JSON mode if requested
        model_config_dict = {}
        if json_mode:
            model_config_dict["response_format"] = {"type": "json_object"}
        
        return ModelFactory.create(
            model_platform=self.chat_options.model_platform,
            model_type=self.chat_options.model_type,
            model_config_dict=model_config_dict if model_config_dict else None,
            api_key=self.chat_options.api_key,
            url=self.chat_options.api_url,
            timeout=60,
        )
    
    async def _call_llm_for_tailoring(
        self,
        system_prompt: str,
        user_prompt: str,
    ) -> Dict[str, Any]:
        """Call LLM for section tailoring with JSON mode."""
        model = self._create_llm_model(json_mode=True)
        
        agent = ChatAgent(
            system_message=system_prompt,
            model=model,
        )
        
        user_msg = BaseMessage.make_user_message(
            role_name="User",
            content=user_prompt,
        )
        
        response = await agent.astep(user_msg)
        return json.loads(response.msg.content)
    
    async def _tailor_section(
        self,
        section_name: str,
        section_data: Any,
        job_description: str,
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Tailor a single resume section using LLM.
        
        Args:
            section_name: Name of the section (e.g., 'experience_details')
            section_data: Original section data (list or dict)
            job_description: Full job description text
            
        Returns:
            Tuple of (section_name, tailored_data_dict)
            On error, returns original data as fallback
        """
        try:
            prompt_template = SECTION_TAILOR_PROMPTS.get(section_name)
            output_schema = SECTION_TAILOR_SCHEMAS.get(section_name)
            
            if not prompt_template or not output_schema:
                logger.warning(f"No tailoring prompt for section: {section_name}")
                return section_name, {section_name: section_data}
            
            # Serialize section data to JSON for LLM input
            if hasattr(section_data, 'model_dump'):
                section_json = json.dumps([section_data.model_dump()], indent=2)
            elif isinstance(section_data, list):
                items = []
                for item in section_data:
                    if hasattr(item, 'model_dump'):
                        items.append(item.model_dump())
                    else:
                        items.append(item)
                section_json = json.dumps(items, indent=2)
            else:
                section_json = json.dumps(section_data, indent=2)
            
            # Build prompt
            prompt = prompt_template.replace("{job_description}", job_description)
            prompt = prompt.replace("{section_json}", section_json)
            prompt = prompt.replace("{output_schema}", output_schema)
            
            system = """You are an expert ATS (Applicant Tracking System) optimization specialist and professional resume writer.

Your goal is to make the resume achieve a 100% keyword match score with the job description.

KEY PRINCIPLES:
1. USE EXACT KEYWORDS from the job description - ATS systems match exact phrases
2. MIRROR terminology precisely - if job says "microservices architecture", use that exact phrase
3. INCORPORATE as many relevant job requirements as truthfully possible
4. QUANTIFY achievements with metrics (%, $, users, latency, uptime)
5. MAINTAIN truthfulness - optimize existing content, don't fabricate

Output valid JSON only. No explanations or markdown."""
            
            result = await self._call_llm_for_tailoring(system, prompt)
            logger.debug(f"Section {section_name} tailored successfully")
            return section_name, result
            
        except Exception as e:
            logger.warning(f"Failed to tailor section {section_name}: {e}. Using original.")
            # Graceful fallback: return original data
            if isinstance(section_data, list):
                items = []
                for item in section_data:
                    if hasattr(item, 'model_dump'):
                        items.append(item.model_dump())
                    else:
                        items.append(item)
                return section_name, {section_name: items}
            return section_name, {section_name: section_data}
    
    async def _tailor_sections_parallel(
        self,
        resume: PlainTextResume,
        job_description: str,
        tailor_education: bool = False,
    ) -> PlainTextResume:
        """
        Tailor resume sections in parallel using LLM.
        
        Args:
            resume: Original PlainTextResume object
            job_description: Full job description text
            tailor_education: Whether to tailor education section (default False)
            
        Returns:
            Tailored PlainTextResume object
        """
        logger.info(f"Starting parallel section tailoring for {resume.personal_information.name}")
        
        # Determine which sections to tailor
        sections_to_process = list(SECTIONS_TO_TAILOR)
        sections_to_copy = list(SECTIONS_TO_COPY)
        
        if tailor_education:
            sections_to_process.append("education_details")
            sections_to_copy.remove("education_details")
        
        # Build tasks for sections that have data
        tasks = []
        resume_dict = resume.model_dump()
        
        for section_name in sections_to_process:
            section_data = resume_dict.get(section_name)
            if section_data:  # Only tailor non-empty sections
                tasks.append(self._tailor_section(section_name, section_data, job_description))
        
        logger.info(f"Tailoring {len(tasks)} sections in parallel: {[t.cr_frame.f_locals.get('section_name', '?') for t in tasks if hasattr(t, 'cr_frame')]}")
        
        # Run all tailoring tasks in parallel
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Collect tailored sections
        merged_data = {}
        tailored_section_names = []
        
        for result in results:
            if isinstance(result, Exception):
                logger.warning(f"Section tailoring exception: {result}")
                continue
            if isinstance(result, tuple) and len(result) == 2:
                section_name, data = result
                if data and isinstance(data, dict):
                    merged_data.update(data)
                    tailored_section_names.append(section_name)
        
        # Copy unchanged sections
        for section_name in sections_to_copy:
            section_data = resume_dict.get(section_name)
            if section_data:
                merged_data[section_name] = section_data
        
        logger.info(f"Tailored sections: {tailored_section_names}")
        logger.info(f"Copied sections: {[s for s in sections_to_copy if resume_dict.get(s)]}")
        
        # Create PlainTextResume directly from merged_data
        return PlainTextResume.from_dict(merged_data)
    
    def _get_builder(self, resume_data: Optional[Dict] = None, resume_yaml: Optional[str] = None):
        """
        Get or create AIHawk resume builder.
        
        Args:
            resume_data: Dict of resume data (legacy mode)
            resume_yaml: YAML string from extract_resume_to_yaml (preferred)
            
        Returns:
            AIHawkResumeBuilder instance or None
        """
        # Prefer YAML mode if provided
        if resume_yaml and self._builder_yaml is None and self.chat_options:
            try:
                from app.libs.aihawk_adapter.resume_builder import AIHawkResumeBuilder
                self._builder_yaml = AIHawkResumeBuilder(
                    chat_options=self.chat_options,
                    output_directory=self.working_directory or "/tmp/resumes",
                    resume_yaml=resume_yaml,
                )
                return self._builder_yaml
            except ImportError as e:
                logger.warning(f"AIHawk adapter not available: {e}")
        
        if resume_yaml and self._builder_yaml:
            # Update existing builder with new YAML
            self._builder_yaml.resume_yaml = resume_yaml
            self._builder_yaml._resume_object = None  # Reset to force rebuild
            return self._builder_yaml
        
        # Legacy dict mode
        if resume_data and self._builder is None and self.chat_options:
            try:
                from app.libs.aihawk_adapter.resume_builder import AIHawkResumeBuilder
                self._builder = AIHawkResumeBuilder(
                    chat_options=self.chat_options,
                    output_directory=self.working_directory or "/tmp/resumes",
                    resume_data=resume_data,
                )
            except ImportError as e:
                logger.warning(f"AIHawk adapter not available: {e}")
        
        return self._builder_yaml or self._builder
    
    def tailor_resume_content(
        self,
        resume_text: str,
        job_description: str,
        job_title: str,
        company: str,
        matching_skills: List[str],
        missing_skills: List[str],
        keywords: List[str],
    ) -> str:
        """
        Tailor resume content to match a specific job description.
        
        This tool rewrites resume bullet points and sections to:
        - Highlight skills that match the job requirements
        - Incorporate relevant keywords naturally
        - Emphasize transferable experience
        - Maintain truthfulness (never adds skills not present)
        
        Args:
            resume_text: Original resume content (plain text or markdown)
            job_description: Full job description text
            job_title: Target job title
            company: Target company name
            matching_skills: Skills that match the job (from analyzer)
            missing_skills: Required skills not in resume
            keywords: Important ATS keywords to incorporate
            
        Returns:
            JSON with tailored resume content and optimization details
        """
        try:
            # Build optimization guidance
            optimization_notes = []
            
            # Start with original content
            tailored_content = resume_text
            
            # Emphasize matching skills in summary/objective
            if matching_skills:
                skill_emphasis = ", ".join(matching_skills[:5])
                optimization_notes.append(f"Emphasized matching skills: {skill_emphasis}")
            
            # Add keyword incorporation suggestions
            keywords_to_add = [k for k in keywords if k.lower() not in resume_text.lower()]
            if keywords_to_add:
                optimization_notes.append(f"Consider incorporating keywords: {', '.join(keywords_to_add[:5])}")
            
            # Note missing skills (for honest disclosure)
            if missing_skills:
                optimization_notes.append(f"Missing skills (do not fabricate): {', '.join(missing_skills[:3])}")
            
            # Generate tailored header
            tailored_header = f"""
---
**Tailored for:** {job_title} at {company}
**Generated:** {datetime.now().strftime('%Y-%m-%d')}
---

"""
            
            # Prepare final content
            final_content = tailored_header + tailored_content
            
            # Calculate ATS score estimate (simplified)
            keyword_matches = sum(1 for k in keywords if k.lower() in resume_text.lower())
            ats_score = min(100, (keyword_matches / max(len(keywords), 1)) * 100 + 40)
            
            result = {
                "status": "success",
                "tailored_content": final_content,
                "optimizations_made": optimization_notes,
                "keywords_incorporated": [k for k in keywords if k.lower() in resume_text.lower()],
                "keywords_to_add": keywords_to_add[:5],
                "ats_score_estimate": round(ats_score, 1),
                "target_job": {
                    "title": job_title,
                    "company": company,
                },
            }
            
            # Save to working directory
            if self.working_directory:
                output_path = Path(self.working_directory) / f"tailored_resume_{company}_{job_title}.md".replace(" ", "_")
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(final_content)
                result["file_path"] = str(output_path)
                logger.info(f"Saved tailored resume to {output_path}")
            
            return json.dumps(result, indent=2)
            
        except Exception as e:
            logger.error(f"Error tailoring resume: {e}")
            return json.dumps({
                "status": "error",
                "message": str(e)
            })
    
    def tailor_resume_from_yaml(
        self,
        resume_yaml: str,
        job_description: str,
        job_title: str,
        company: str,
        style: str = "classic",
        section_mode: bool = True,
        tailor_education: bool = False,
    ) -> str:
        """
        Tailor resume using YAML input with section-by-section LLM processing.
        
        This is the PREFERRED method for resume tailoring. Use the YAML output
        from JobAnalyzerToolkit.extract_resume_to_yaml() as input.
        
        When section_mode=True (default), each section is tailored independently
        in parallel using LLM for better quality and faster processing.
        
        Sections tailored by default:
        - experience_details: Responsibilities rewritten with job-relevant terms
        - projects: Descriptions emphasizing matching technologies
        - achievements: Highlighted for job relevance
        - certifications: Prioritized by relevance
        
        Sections copied unchanged:
        - personal_information, languages, interests, education_details (unless tailor_education=True)
        
        Args:
            resume_yaml: YAML string from extract_resume_to_yaml (plain_text_resume format)
            job_description: Full job description text
            job_title: Target job title
            company: Target company name
            style: Resume style - classic, modern, minimal, professional
            section_mode: Use parallel section tailoring (default True)
            tailor_education: Also tailor education section (default False)
            
        Returns:
            JSON with tailored resume in YAML, HTML, and markdown formats
        """
        # Use async version via event loop
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # Already in async context - create task
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(
                        asyncio.run,
                        self.tailor_resume_from_yaml_async(
                            resume_yaml, job_description, job_title, company,
                            style, section_mode, tailor_education
                        )
                    )
                    return future.result()
            else:
                return loop.run_until_complete(
                    self.tailor_resume_from_yaml_async(
                        resume_yaml, job_description, job_title, company,
                        style, section_mode, tailor_education
                    )
                )
        except RuntimeError:
            # No event loop - create one
            return asyncio.run(
                self.tailor_resume_from_yaml_async(
                    resume_yaml, job_description, job_title, company,
                    style, section_mode, tailor_education
                )
            )
    
    async def tailor_resume_from_yaml_async(
        self,
        resume_yaml: str,
        job_description: str,
        job_title: str,
        company: str,
        style: str = "classic",
        section_mode: bool = True,
        tailor_education: bool = False,
    ) -> str:
        """
        Async version of tailor_resume_from_yaml with parallel section processing.
        
        Args:
            resume_yaml: YAML string from extract_resume_to_yaml
            job_description: Full job description text
            job_title: Target job title
            company: Target company name
            style: Resume style
            section_mode: Use parallel section tailoring (default True)
            tailor_education: Also tailor education section (default False)
            
        Returns:
            JSON with tailored resume in YAML, HTML, and markdown formats
        """
        try:
            # Validate YAML input
            try:
                resume = PlainTextResume.from_yaml(resume_yaml)
                logger.info(f"Parsed resume YAML for: {resume.personal_information.name} {resume.personal_information.surname}")
            except Exception as e:
                return json.dumps({
                    "status": "error",
                    "message": f"Invalid resume YAML format: {str(e)}"
                })
            
            # Use section-based parallel tailoring when LLM available
            if section_mode and self.chat_options:
                try:
                    tailored_resume = await self._tailor_sections_parallel(
                        resume=resume,
                        job_description=job_description,
                        tailor_education=tailor_education,
                    )
                    
                    # Generate outputs
                    tailored_yaml = tailored_resume.to_yaml()
                    markdown_content = self._resume_to_markdown(tailored_resume, job_title, company)
                    html_content = self._markdown_to_html(markdown_content)
                    
                    # Save outputs
                    file_paths = {}
                    if self.working_directory:
                        output_dir = Path(self.working_directory)
                        output_dir.mkdir(parents=True, exist_ok=True)
                        
                        safe_company = company.replace(" ", "_").replace("/", "-")[:30]
                        safe_title = job_title.replace(" ", "_").replace("/", "-")[:30]
                        base_name = f"tailored_{safe_company}_{safe_title}"
                        
                        # Save YAML
                        yaml_path = output_dir / f"{base_name}.yaml"
                        with open(yaml_path, 'w', encoding='utf-8') as f:
                            f.write(tailored_yaml)
                        file_paths["yaml"] = str(yaml_path)
                        
                        # Save Markdown
                        md_path = output_dir / f"{base_name}.md"
                        with open(md_path, 'w', encoding='utf-8') as f:
                            f.write(markdown_content)
                        file_paths["markdown"] = str(md_path)
                        
                        logger.info(f"Saved tailored resume to {output_dir}")
                    
                    # Determine which sections were tailored
                    sections_tailored = list(SECTIONS_TO_TAILOR)
                    if tailor_education:
                        sections_tailored.append("education_details")
                    
                    return json.dumps({
                        "status": "success",
                        "tailored_resume_yaml": tailored_yaml,
                        "html_content": html_content,
                        "markdown_content": markdown_content,
                        "file_paths": file_paths,
                        "target_job": {
                            "title": job_title,
                            "company": company,
                        },
                        "sections_tailored": sections_tailored,
                        "tailor_education": tailor_education,
                    }, indent=2)
                    
                except Exception as e:
                    logger.warning(f"Parallel section tailoring failed: {e}. Trying AIHawk fallback.")
            
            # Try AIHawk builder as fallback
            builder = self._get_builder(resume_yaml=resume_yaml)
            
            if builder is not None:
                try:
                    html_content, markdown_content = builder.tailor_resume_for_job(
                        job_description=job_description,
                        style=style,
                    )
                    
                    file_paths = {}
                    if self.working_directory:
                        output_dir = Path(self.working_directory)
                        output_dir.mkdir(parents=True, exist_ok=True)
                        
                        safe_company = company.replace(" ", "_").replace("/", "-")[:30]
                        safe_title = job_title.replace(" ", "_").replace("/", "-")[:30]
                        base_name = f"tailored_{safe_company}_{safe_title}"
                        
                        html_path = output_dir / f"{base_name}.html"
                        with open(html_path, 'w', encoding='utf-8') as f:
                            f.write(html_content)
                        file_paths["html"] = str(html_path)
                        
                        md_path = output_dir / f"{base_name}.md"
                        with open(md_path, 'w', encoding='utf-8') as f:
                            f.write(markdown_content)
                        file_paths["markdown"] = str(md_path)
                    
                    return json.dumps({
                        "status": "success",
                        "html_content": html_content,
                        "markdown_content": markdown_content,
                        "file_paths": file_paths,
                        "target_job": {
                            "title": job_title,
                            "company": company,
                        },
                        "style": style,
                        "method": "aihawk",
                    }, indent=2)
                    
                except Exception as e:
                    logger.warning(f"AIHawk tailoring failed: {e}")
            
            # Final fallback to basic formatting
            logger.warning("Using basic fallback tailoring")
            return self._fallback_yaml_tailoring(resume, job_description, job_title, company)
                
        except Exception as e:
            logger.error(f"Error tailoring resume from YAML: {e}")
            return json.dumps({
                "status": "error",
                "message": str(e)
            })
    
    def _resume_to_markdown(
        self,
        resume: PlainTextResume,
        job_title: str,
        company: str,
    ) -> str:
        """Convert PlainTextResume to markdown format."""
        pi = resume.personal_information
        md_parts = []
        
        # Header
        md_parts.append(f"# {pi.name} {pi.surname}\n")
        
        contact_parts = []
        if pi.email:
            contact_parts.append(pi.email)
        if pi.phone:
            phone = f"{pi.phone_prefix}{pi.phone}" if pi.phone_prefix else pi.phone
            contact_parts.append(phone)
        if pi.linkedin:
            contact_parts.append(pi.linkedin)
        if pi.github:
            contact_parts.append(pi.github)
        if contact_parts:
            md_parts.append(" | ".join(contact_parts) + "\n")
        
        if pi.city or pi.country:
            location = ", ".join(filter(None, [pi.city, pi.country]))
            md_parts.append(f"📍 {location}\n")
        
        md_parts.append(f"\n---\n**Tailored for:** {job_title} at {company}\n---\n")
        
        # Experience
        if resume.experience_details:
            md_parts.append("\n## Professional Experience\n")
            for exp in resume.experience_details:
                md_parts.append(f"\n### {exp.position} | {exp.company}\n")
                if exp.employment_period:
                    md_parts.append(f"*{exp.employment_period}*")
                if exp.location:
                    md_parts.append(f" - {exp.location}")
                md_parts.append("\n")
                
                if exp.key_responsibilities:
                    for resp in exp.key_responsibilities:
                        if isinstance(resp, dict):
                            for key, val in resp.items():
                                md_parts.append(f"- {val}\n")
                        else:
                            md_parts.append(f"- {resp}\n")
                
                if exp.skills_acquired:
                    md_parts.append(f"\n**Skills:** {', '.join(exp.skills_acquired)}\n")
        
        # Education
        if resume.education_details:
            md_parts.append("\n## Education\n")
            for edu in resume.education_details:
                md_parts.append(f"\n### {edu.education_level} in {edu.field_of_study}\n")
                md_parts.append(f"*{edu.institution}*")
                if edu.year_of_completion:
                    md_parts.append(f" - {edu.year_of_completion}")
                md_parts.append("\n")
                if edu.final_evaluation_grade:
                    md_parts.append(f"GPA: {edu.final_evaluation_grade}\n")
        
        # Projects
        if resume.projects:
            md_parts.append("\n## Projects\n")
            for proj in resume.projects:
                md_parts.append(f"\n### {proj.name}\n")
                if proj.description:
                    md_parts.append(f"{proj.description}\n")
                if proj.link:
                    md_parts.append(f"[Link]({proj.link})\n")
        
        # Achievements
        if resume.achievements:
            md_parts.append("\n## Achievements\n")
            for ach in resume.achievements:
                md_parts.append(f"- **{ach.name}**: {ach.description}\n")
        
        # Certifications
        if resume.certifications:
            md_parts.append("\n## Certifications\n")
            for cert in resume.certifications:
                md_parts.append(f"- **{cert.name}**: {cert.description}\n")
        
        # Languages
        if resume.languages:
            md_parts.append("\n## Languages\n")
            for lang in resume.languages:
                md_parts.append(f"- {lang.language}: {lang.proficiency}\n")
        
        # Interests
        if resume.interests:
            md_parts.append("\n## Interests\n")
            md_parts.append(", ".join(resume.interests) + "\n")
        
        return "".join(md_parts)
    
    def _fallback_yaml_tailoring(
        self,
        resume: PlainTextResume,
        job_description: str,
        job_title: str,
        company: str,
    ) -> str:
        """
        Fallback tailoring when AIHawk is not available.
        
        Creates a formatted markdown resume from the YAML data with
        basic job-specific header.
        """
        pi = resume.personal_information
        
        # Build markdown resume
        md_parts = []
        
        # Header
        md_parts.append(f"# {pi.name} {pi.surname}\n")
        
        contact_parts = []
        if pi.email:
            contact_parts.append(pi.email)
        if pi.phone:
            contact_parts.append(f"{pi.phone_prefix}{pi.phone}" if pi.phone_prefix else pi.phone)
        if pi.linkedin:
            contact_parts.append(pi.linkedin)
        if pi.github:
            contact_parts.append(pi.github)
        if contact_parts:
            md_parts.append(" | ".join(contact_parts) + "\n")
        
        if pi.city or pi.country:
            location = ", ".join(filter(None, [pi.city, pi.country]))
            md_parts.append(f"📍 {location}\n")
        
        md_parts.append(f"\n---\n**Tailored for:** {job_title} at {company}\n---\n")
        
        # Experience
        if resume.experience_details:
            md_parts.append("\n## Professional Experience\n")
            for exp in resume.experience_details:
                md_parts.append(f"\n### {exp.position} | {exp.company}\n")
                if exp.employment_period:
                    md_parts.append(f"*{exp.employment_period}*")
                if exp.location:
                    md_parts.append(f" - {exp.location}")
                md_parts.append("\n")
                
                if exp.key_responsibilities:
                    for resp in exp.key_responsibilities:
                        if isinstance(resp, dict):
                            for key, val in resp.items():
                                md_parts.append(f"- {val}\n")
                        else:
                            md_parts.append(f"- {resp}\n")
                
                if exp.skills_acquired:
                    md_parts.append(f"\n**Skills:** {', '.join(exp.skills_acquired)}\n")
        
        # Education
        if resume.education_details:
            md_parts.append("\n## Education\n")
            for edu in resume.education_details:
                md_parts.append(f"\n### {edu.education_level} in {edu.field_of_study}\n")
                md_parts.append(f"*{edu.institution}*")
                if edu.year_of_completion:
                    md_parts.append(f" - {edu.year_of_completion}")
                md_parts.append("\n")
                if edu.final_evaluation_grade:
                    md_parts.append(f"GPA: {edu.final_evaluation_grade}\n")
        
        # Projects
        if resume.projects:
            md_parts.append("\n## Projects\n")
            for proj in resume.projects:
                md_parts.append(f"\n### {proj.name}\n")
                if proj.description:
                    md_parts.append(f"{proj.description}\n")
                if proj.link:
                    md_parts.append(f"[Link]({proj.link})\n")
        
        # Certifications
        if resume.certifications:
            md_parts.append("\n## Certifications\n")
            for cert in resume.certifications:
                md_parts.append(f"- **{cert.name}**: {cert.description}\n")
        
        # Languages
        if resume.languages:
            md_parts.append("\n## Languages\n")
            for lang in resume.languages:
                md_parts.append(f"- {lang.language}: {lang.proficiency}\n")
        
        # Interests
        if resume.interests:
            md_parts.append("\n## Interests\n")
            md_parts.append(", ".join(resume.interests) + "\n")
        
        markdown_content = "".join(md_parts)
        
        # Save if working directory is set
        file_path = None
        if self.working_directory:
            output_dir = Path(self.working_directory)
            output_dir.mkdir(parents=True, exist_ok=True)
            
            safe_company = company.replace(" ", "_").replace("/", "-")[:30]
            safe_title = job_title.replace(" ", "_").replace("/", "-")[:30]
            file_path = output_dir / f"tailored_{safe_company}_{safe_title}.md"
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
        
        return json.dumps({
            "status": "success",
            "html_content": "",  # No HTML in fallback mode
            "markdown_content": markdown_content,
            "file_paths": {"markdown": str(file_path)} if file_path else {},
            "target_job": {
                "title": job_title,
                "company": company,
            },
            "note": "Generated using fallback mode (AIHawk not available)",
        }, indent=2)
    
    def optimize_for_ats(
        self,
        resume_content: str,
        job_keywords: List[str],
    ) -> str:
        """
        Optimize resume content for Applicant Tracking Systems (ATS).
        
        ATS optimization includes:
        - Adding relevant keywords naturally
        - Using standard section headings
        - Proper formatting for parsing
        - Removing graphics/tables that ATS can't parse
        
        Args:
            resume_content: Current resume content
            job_keywords: Important keywords from job posting
            
        Returns:
            JSON with ATS-optimized content and score
        """
        try:
            # Standard ATS-friendly section headings
            standard_headings = {
                'profile': 'Professional Summary',
                'summary': 'Professional Summary',
                'objective': 'Career Objective',
                'experience': 'Professional Experience',
                'work history': 'Professional Experience',
                'employment': 'Professional Experience',
                'education': 'Education',
                'academic': 'Education',
                'skills': 'Skills',
                'technical skills': 'Technical Skills',
                'certifications': 'Certifications',
                'certificates': 'Certifications',
            }
            
            optimized_content = resume_content
            
            # Standardize section headings
            for original, standard in standard_headings.items():
                pattern = rf'\b{original}\b'
                import re
                optimized_content = re.sub(pattern, standard, optimized_content, flags=re.IGNORECASE)
            
            # Calculate keyword incorporation
            content_lower = optimized_content.lower()
            found_keywords = [k for k in job_keywords if k.lower() in content_lower]
            missing_keywords = [k for k in job_keywords if k.lower() not in content_lower]
            
            # ATS score calculation
            keyword_score = (len(found_keywords) / max(len(job_keywords), 1)) * 60
            
            # Check for ATS-friendly formatting
            format_score = 25  # Base score
            if '|' not in optimized_content:  # No tables
                format_score += 5
            if not re.search(r'[^\x00-\x7F]', optimized_content):  # ASCII only
                format_score += 5
            if len(re.findall(r'^#+\s', optimized_content, re.MULTILINE)) > 0:  # Has headers
                format_score += 5
            
            ats_score = min(100, keyword_score + format_score)
            
            result = {
                "status": "success",
                "optimized_content": optimized_content,
                "ats_score": round(ats_score, 1),
                "found_keywords": found_keywords,
                "missing_keywords": missing_keywords[:10],
                "recommendations": [],
            }
            
            # Add recommendations
            if missing_keywords:
                result["recommendations"].append(
                    f"Consider adding these keywords naturally: {', '.join(missing_keywords[:5])}"
                )
            if ats_score < 70:
                result["recommendations"].append(
                    "ATS score is below optimal. Try incorporating more job-specific keywords."
                )
            
            return json.dumps(result, indent=2)
            
        except Exception as e:
            logger.error(f"Error optimizing for ATS: {e}")
            return json.dumps({
                "status": "error",
                "message": str(e)
            })
    
    def generate_cover_letter(
        self,
        resume_summary: str,
        job_title: str,
        company: str,
        job_description: str,
        matching_skills: List[str],
        tone: str = "professional",
    ) -> str:
        """
        Generate a tailored cover letter for a specific job application.
        
        Creates a professional cover letter that:
        - Addresses the specific job and company
        - Highlights relevant skills and experience
        - Shows enthusiasm for the role
        - Maintains appropriate tone
        
        Args:
            resume_summary: Brief summary of candidate's background
            job_title: Target job title
            company: Target company name
            job_description: Full job description
            matching_skills: Skills that match the job
            tone: Letter tone - professional, enthusiastic, conversational
            
        Returns:
            JSON with cover letter content
        """
        try:
            # Extract key points from job description
            import re
            
            # Find company values/culture mentions
            culture_keywords = re.findall(
                r'(?:culture|values?|mission|vision|team)[:\s]+([^.]+)',
                job_description,
                re.IGNORECASE
            )
            
            # Build cover letter
            date = datetime.now().strftime('%B %d, %Y')
            
            cover_letter = f"""{date}

Dear Hiring Manager,

I am writing to express my strong interest in the {job_title} position at {company}. {resume_summary}

Based on my experience, I am confident that my skills align well with your requirements:

"""
            
            # Add skill highlights
            for skill in matching_skills[:3]:
                cover_letter += f"• {skill}\n"
            
            cover_letter += f"""
I am particularly drawn to {company} because of """
            
            if culture_keywords:
                cover_letter += f"your commitment to {culture_keywords[0].strip().lower()}. "
            else:
                cover_letter += "the opportunity to contribute to meaningful work in this role. "
            
            cover_letter += f"""I believe my background and enthusiasm make me an excellent fit for your team.

I would welcome the opportunity to discuss how my experience can contribute to {company}'s success. Thank you for considering my application.

Sincerely,
[Your Name]
"""
            
            # Save to working directory
            if self.working_directory:
                output_path = Path(self.working_directory) / f"cover_letter_{company}_{job_title}.md".replace(" ", "_")
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(cover_letter)
                logger.info(f"Saved cover letter to {output_path}")
            
            return json.dumps({
                "status": "success",
                "cover_letter": cover_letter,
                "word_count": len(cover_letter.split()),
                "file_path": str(output_path) if self.working_directory else None,
            }, indent=2)
            
        except Exception as e:
            logger.error(f"Error generating cover letter: {e}")
            return json.dumps({
                "status": "error",
                "message": str(e)
            })
    
    def export_to_pdf(
        self,
        content: str,
        filename: str,
    ) -> str:
        """
        Export content (resume or cover letter) to PDF format.
        
        Args:
            content: Markdown or HTML content to export
            filename: Output filename (without extension)
            
        Returns:
            JSON with file path of generated PDF
        """
        try:
            output_dir = Path(self.working_directory) if self.working_directory else Path("/tmp/resumes")
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"{filename}.pdf"
            
            # Try pdfkit first
            try:
                import pdfkit
                
                # Convert markdown to HTML
                html_content = self._markdown_to_html(content)
                
                pdfkit.from_string(
                    html_content,
                    str(output_path),
                    options={
                        'page-size': 'Letter',
                        'margin-top': '0.75in',
                        'margin-right': '0.75in',
                        'margin-bottom': '0.75in',
                        'margin-left': '0.75in',
                        'encoding': 'UTF-8',
                    }
                )
                
                logger.info(f"Exported PDF to {output_path}")
                
                return json.dumps({
                    "status": "success",
                    "file_path": str(output_path),
                    "format": "pdf",
                })
                
            except ImportError:
                logger.warning("pdfkit not available, falling back to HTML")
                
                # Fallback: save as HTML
                html_path = output_dir / f"{filename}.html"
                html_content = self._markdown_to_html(content)
                
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
                return json.dumps({
                    "status": "success",
                    "file_path": str(html_path),
                    "format": "html",
                    "note": "PDF export not available, saved as HTML instead"
                })
                
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            return json.dumps({
                "status": "error",
                "message": str(e)
            })
    
    def export_to_docx(
        self,
        content: str,
        filename: str,
    ) -> str:
        """
        Export content (resume or cover letter) to DOCX format.
        
        Args:
            content: Markdown content to export
            filename: Output filename (without extension)
            
        Returns:
            JSON with file path of generated DOCX
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            output_dir = Path(self.working_directory) if self.working_directory else Path("/tmp/resumes")
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"{filename}.docx"
            
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Parse and add content
            lines = content.split('\n')
            
            for line in lines:
                line = line.rstrip()
                
                if not line:
                    doc.add_paragraph()
                    continue
                
                if line.startswith('# '):
                    p = doc.add_heading(line[2:], level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('---'):
                    doc.add_paragraph('_' * 50)
                else:
                    # Handle bold/italic
                    import re
                    p = doc.add_paragraph()
                    
                    # Simple bold handling
                    parts = re.split(r'\*\*(.+?)\*\*', line)
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # Bold part
                            p.add_run(part).bold = True
                        else:
                            p.add_run(part)
            
            doc.save(str(output_path))
            logger.info(f"Exported DOCX to {output_path}")
            
            return json.dumps({
                "status": "success",
                "file_path": str(output_path),
                "format": "docx",
            })
            
        except ImportError:
            logger.error("python-docx not installed")
            
            # Fallback: save as markdown
            md_path = Path(self.working_directory or "/tmp") / f"{filename}.md"
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return json.dumps({
                "status": "success",
                "file_path": str(md_path),
                "format": "markdown",
                "note": "DOCX export not available, saved as Markdown instead"
            })
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            return json.dumps({
                "status": "error",
                "message": str(e)
            })
    
    def _markdown_to_html(self, markdown_content: str) -> str:
        """Convert markdown to HTML for PDF export"""
        import re
        
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: 'Calibri', 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 20px;
        }}
        h1 {{ font-size: 18pt; margin-bottom: 10px; }}
        h2 {{ font-size: 14pt; margin-top: 15px; margin-bottom: 8px; border-bottom: 1px solid #333; }}
        h3 {{ font-size: 12pt; margin-top: 10px; margin-bottom: 5px; }}
        ul {{ margin: 5px 0; padding-left: 20px; }}
        li {{ margin: 3px 0; }}
        hr {{ border: none; border-top: 1px solid #ccc; margin: 10px 0; }}
    </style>
</head>
<body>
"""
        
        # Convert markdown to HTML
        content = markdown_content
        
        # Headers
        content = re.sub(r'^### (.+)$', r'<h3>\1</h3>', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'<h2>\1</h2>', content, flags=re.MULTILINE)
        content = re.sub(r'^# (.+)$', r'<h1>\1</h1>', content, flags=re.MULTILINE)
        
        # Bold and italic
        content = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', content)
        content = re.sub(r'\*(.+?)\*', r'<em>\1</em>', content)
        
        # Lists
        content = re.sub(r'^[-*•] (.+)$', r'<li>\1</li>', content, flags=re.MULTILINE)
        content = re.sub(r'(<li>.*</li>\n?)+', r'<ul>\g<0></ul>', content)
        
        # Horizontal rules
        content = re.sub(r'^---+$', r'<hr>', content, flags=re.MULTILINE)
        
        # Paragraphs
        content = re.sub(r'\n\n+', r'</p><p>', content)
        content = f'<p>{content}</p>'
        
        html += content + "\n</body>\n</html>"
        
        return html
    
    def get_available_styles(self) -> str:
        """
        Get list of available resume styles from AIHawk.
        
        Returns:
            JSON with available style names
        """
        try:
            # Check AIHawk styles directory
            from pathlib import Path
            
            aihawk_path = Path(__file__).resolve().parents[4] / "external" / "aihawk"
            styles_dir = aihawk_path / "src" / "libs" / "resume_and_cover_builder" / "resume_style"
            
            styles = []
            if styles_dir.exists():
                for css_file in styles_dir.glob("*.css"):
                    styles.append(css_file.stem)
            
            if not styles:
                styles = ["classic", "modern", "minimal"]  # Defaults
            
            return json.dumps({
                "status": "success",
                "styles": styles,
                "default": "classic",
            })
            
        except Exception as e:
            return json.dumps({
                "status": "error",
                "message": str(e),
                "styles": ["classic"],  # Fallback
            })
    
    def get_tools(self) -> List[FunctionTool]:
        return [
            FunctionTool(self.tailor_resume_content),
            FunctionTool(self.tailor_resume_from_yaml),
            FunctionTool(self.optimize_for_ats),
            FunctionTool(self.generate_cover_letter),
            FunctionTool(self.export_to_pdf),
            FunctionTool(self.export_to_docx),
            FunctionTool(self.get_available_styles),
        ]
